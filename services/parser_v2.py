"""
Vision-first парсер — все файлы через Vision API
"""
import io
import base64
from pathlib import Path
from typing import Optional, List
import fitz  # PyMuPDF
import pandas as pd
from PIL import Image
from openai import AsyncOpenAI

from config import OPENAI_API_KEY

client = AsyncOpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None

# Промпт для Vision — извлечение ВСЕГО
VISION_EXTRACT_PROMPT = """Ты анализируешь документ о жилом комплексе для риэлтора.

Извлеки ВСЮ информацию с изображения:

1. ТЕКСТ — перепиши весь текст как есть
2. ЦЕНЫ — если есть цены, выпиши каждую отдельно (21 172 800 ₽, 33 680 640 ₽, и т.д.)
3. ПЛОЩАДИ — все площади квартир (24.06 м², 42.96 м², и т.д.)
4. ТАБЛИЦЫ — если есть таблица, сохрани структуру
5. ПЛАНИРОВКИ — опиши: тип квартиры, площадь, комнаты, особенности
6. УСЛОВИЯ — рассрочка, ипотека, первый взнос, сроки

ВАЖНО:
- Не пропускай числа и цены
- Извлекай ВСЕ данные, даже мелкий текст
- Если видишь несколько цен — выпиши все
- Пиши на русском"""


async def extract_all(file_path: str) -> str:
    """Главная функция — извлечь всё из файла"""
    path = Path(file_path)
    suffix = path.suffix.lower()
    
    try:
        if suffix == ".pdf":
            return await extract_pdf_vision(file_path)
        
        elif suffix in (".jpg", ".jpeg", ".png", ".webp"):
            return await extract_image_vision(file_path)
        
        elif suffix == ".docx":
            return extract_from_docx(file_path)
        
        elif suffix in (".xlsx", ".xls"):
            return extract_from_excel(file_path)
        
        elif suffix == ".csv":
            return extract_from_csv(file_path)
        
        elif suffix == ".txt":
            return path.read_text(encoding="utf-8", errors="ignore")
        
        else:
            return f"[Неподдерживаемый формат: {suffix}]"
            
    except Exception as e:
        print(f"[PARSER_V2] Error: {file_path} — {e}")
        return f"[Ошибка: {e}]"


async def extract_pdf_vision(file_path: str, max_pages: int = 30) -> str:
    """PDF → изображения страниц → Vision API"""
    
    if not client:
        return "[OpenAI не настроен]"
    
    images_base64 = []
    
    try:
        with fitz.open(file_path) as doc:
            num_pages = min(len(doc), max_pages)
            print(f"[PARSER_V2] PDF {Path(file_path).name}: {num_pages} страниц")
            
            for page_num in range(num_pages):
                page = doc[page_num]
                # Рендерим страницу в изображение (150 DPI для баланса качество/размер)
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("jpeg")
                img_base64 = base64.b64encode(img_bytes).decode()
                images_base64.append((page_num + 1, img_base64))
                
    except Exception as e:
        print(f"[PARSER_V2] PDF render error: {e}")
        return f"[Ошибка рендеринга PDF: {e}]"
    
    if not images_base64:
        return "[PDF пустой]"
    
    # Отправляем все страницы в Vision API
    all_text = []
    
    for page_num, img_base64 in images_base64:
        try:
            text = await _call_vision_api(img_base64)
            if text and not text.startswith("["):
                all_text.append(f"=== СТРАНИЦА {page_num} ===\n{text}")
                print(f"[PARSER_V2] Страница {page_num}: {len(text)} символов")
        except Exception as e:
            print(f"[PARSER_V2] Vision error page {page_num}: {e}")
            all_text.append(f"=== СТРАНИЦА {page_num} ===\n[Ошибка распознавания]")
    
    return "\n\n".join(all_text)


async def extract_image_vision(file_path: str) -> str:
    """Изображение → Vision API"""
    
    if not client:
        return "[OpenAI не настроен]"
    
    try:
        with Image.open(file_path) as img:
            # Конвертируем в RGB если нужно
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            
            # Уменьшаем если слишком большое
            max_size = 2000
            if max(img.size) > max_size:
                ratio = max_size / max(img.size)
                new_size = (int(img.size[0] * ratio), int(img.size[1] * ratio))
                img = img.resize(new_size, Image.Resampling.LANCZOS)
            
            buffer = io.BytesIO()
            img.save(buffer, format="JPEG", quality=85)
            img_base64 = base64.b64encode(buffer.getvalue()).decode()
        
        text = await _call_vision_api(img_base64)
        print(f"[PARSER_V2] Image {Path(file_path).name}: {len(text)} символов")
        return text
        
    except Exception as e:
        print(f"[PARSER_V2] Image error: {e}")
        return f"[Ошибка изображения: {e}]"


async def _call_vision_api(image_base64: str) -> str:
    """Вызов Vision API для одного изображения"""
    
    try:
        response = await client.chat.completions.create(
            model="gpt-4o",  # Используем gpt-4o для лучшего качества Vision
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": VISION_EXTRACT_PROMPT},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{image_base64}",
                                "detail": "high"  # Высокое качество распознавания
                            }
                        }
                    ]
                }
            ],
            max_tokens=2000,
            temperature=0.1
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        print(f"[PARSER_V2] Vision API error: {e}")
        return f"[Ошибка Vision API: {e}]"


def extract_from_docx(file_path: str) -> str:
    """DOCX — текст + таблицы"""
    try:
        from docx import Document
        doc = Document(file_path)
        
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_rows.append(" | ".join(cells))
            if table_rows:
                text_parts.append("\n[ТАБЛИЦА]\n" + "\n".join(table_rows))
        
        return "\n".join(text_parts) if text_parts else "[DOCX пустой]"
        
    except Exception as e:
        return f"[Ошибка DOCX: {e}]"


def extract_from_excel(file_path: str) -> str:
    """Excel — все листы"""
    text_parts = []
    
    try:
        xlsx = pd.ExcelFile(file_path)
        
        for sheet_name in xlsx.sheet_names:
            df = pd.read_excel(xlsx, sheet_name=sheet_name)
            df = df.dropna(how='all').dropna(axis=1, how='all')
            
            if df.empty:
                continue
            
            df = df.fillna('')
            text_parts.append(f"=== ЛИСТ: {sheet_name} ===")
            
            for idx, row in df.iterrows():
                row_values = [str(v).strip() for v in row.values if str(v).strip()]
                if row_values:
                    text_parts.append(" | ".join(row_values))
        
    except Exception as e:
        return f"[Ошибка Excel: {e}]"
    
    return "\n".join(text_parts) if text_parts else "[Excel пустой]"


def extract_from_csv(file_path: str) -> str:
    """CSV"""
    try:
        try:
            df = pd.read_csv(file_path, encoding="utf-8")
        except:
            df = pd.read_csv(file_path, encoding="cp1251")
        
        if df.empty:
            return "[CSV пустой]"
        
        return df.to_string(index=False)
        
    except Exception as e:
        return f"[Ошибка CSV: {e}]"


def get_file_info(file_path: str) -> dict:
    """Информация о файле"""
    path = Path(file_path)
    return {
        "name": path.name,
        "extension": path.suffix.lower(),
        "size_kb": path.stat().st_size // 1024 if path.exists() else 0,
        "exists": path.exists()
    }
