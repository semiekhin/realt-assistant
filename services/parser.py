"""
Сервис для извлечения текста из файлов
"""
import io
import base64
from pathlib import Path
from typing import Optional
import fitz  # PyMuPDF
import pandas as pd
from PIL import Image

from services.llm import extract_text_from_image


async def extract_text(file_path: str) -> str:
    """Извлечь текст из файла"""
    path = Path(file_path)
    suffix = path.suffix.lower()
    
    try:
        if suffix == ".pdf":
            return extract_from_pdf(file_path)
        
        elif suffix in (".xlsx", ".xls"):
            return extract_from_excel(file_path)
        
        elif suffix == ".csv":
            return extract_from_csv(file_path)
        
        elif suffix in (".jpg", ".jpeg", ".png", ".webp"):
            return await extract_from_image(file_path)
        
        elif suffix == ".txt":
            return path.read_text(encoding="utf-8", errors="ignore")
        
        elif suffix == ".docx":
            return extract_from_docx(file_path)
        
        else:
            return f"[Неподдерживаемый формат: {suffix}]"
            
    except Exception as e:
        print(f"[PARSER] Error extracting from {file_path}: {e}")
        return f"[Ошибка чтения файла: {e}]"


def extract_from_pdf(file_path: str) -> str:
    """Извлечь текст из PDF"""
    text_parts = []
    
    try:
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc, 1):
                text = page.get_text().strip()
                if text and len(text) > 20:
                    text_parts.append(f"--- Страница {page_num} ---\n{text}")
    except Exception as e:
        print(f"[PARSER] PDF error: {e}")
        return f"[Ошибка чтения PDF: {e}]"
    
    if not text_parts:
        return "[PDF без текста - возможно содержит только изображения]"
    
    return "\n\n".join(text_parts)


def extract_from_excel(file_path: str) -> str:
    """Извлечь данные из Excel"""
    text_parts = []
    
    try:
        xlsx = pd.ExcelFile(file_path)
        
        for sheet_name in xlsx.sheet_names:
            df = pd.read_excel(xlsx, sheet_name=sheet_name)
            
            if df.empty:
                continue
            
            # Убираем полностью пустые строки и столбцы
            df = df.dropna(how='all').dropna(axis=1, how='all')
            
            if df.empty:
                continue
            
            # Заменяем NaN на пустую строку
            df = df.fillna('')
            
            text_parts.append(f"--- Лист: {sheet_name} ---")
            
            # Форматируем красиво
            for idx, row in df.iterrows():
                row_values = [str(v).strip() for v in row.values if str(v).strip()]
                if row_values:
                    text_parts.append(" | ".join(row_values))
        
    except Exception as e:
        print(f"[PARSER] Excel error: {e}")
        return f"[Ошибка чтения Excel: {e}]"
    
    return "\n".join(text_parts) if text_parts else "[Excel пустой]"


def extract_from_csv(file_path: str) -> str:
    """Извлечь данные из CSV"""
    try:
        try:
            df = pd.read_csv(file_path, encoding="utf-8")
        except:
            df = pd.read_csv(file_path, encoding="cp1251")
        
        if df.empty:
            return "[CSV пустой]"
        
        df = df.fillna('')
        return df.to_string(index=False)
        
    except Exception as e:
        return f"[Ошибка CSV: {e}]"


def extract_from_docx(file_path: str) -> str:
    """Извлечь текст из DOCX"""
    try:
        from docx import Document
        doc = Document(file_path)
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    text_parts.append(" | ".join(cells))
        
        return "\n".join(text_parts) if text_parts else "[DOCX пустой]"
        
    except Exception as e:
        return f"[Ошибка DOCX: {e}]"


async def extract_from_image(file_path: str) -> str:
    """Извлечь текст из изображения через Vision API"""
    
    with Image.open(file_path) as img:
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        
        max_size = 2000
        if max(img.size) > max_size:
            ratio = max_size / max(img.size)
            new_size = (int(img.size[0] * ratio), int(img.size[1] * ratio))
            img = img.resize(new_size, Image.Resampling.LANCZOS)
        
        buffer = io.BytesIO()
        img.save(buffer, format="JPEG", quality=85)
        image_base64 = base64.b64encode(buffer.getvalue()).decode()
    
    text = await extract_text_from_image(image_base64)
    
    return text if text else "[Не удалось распознать текст]"


def get_file_info(file_path: str) -> dict:
    """Получить информацию о файле"""
    path = Path(file_path)
    
    return {
        "name": path.name,
        "extension": path.suffix.lower(),
        "size_kb": path.stat().st_size // 1024 if path.exists() else 0,
        "exists": path.exists()
    }
